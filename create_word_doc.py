"""
create_word_doc.py — Generate a professional Microsoft Word (.docx) setup guide for colleagues.
Run: python create_word_doc.py
"""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

OUTPUT_DOCX = "Kezzler_Mock_Database_Setup_Guide.docx"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_guide():
    doc = Document()

    # Page Margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Kezzler Mock Database & AI Setup Guide")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 37, 66)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Step-by-Step Instructions for Importing the 51-Table Mock Database & Running the Agent")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spacing

    # Section 1: Overview
    h1 = doc.add_heading("1. Overview & Contents", level=1)
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    p = doc.add_paragraph()
    p.add_run("This package contains the full 51-table MySQL database structure for Kezzler, populated with 632 records of realistic Track-and-Trace agribusiness data (products, distributors, sales transactions, field sales reps, QR codes, and anti-counterfeit logs).")

    # Table of Files
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "File Name"
    hdr[1].text = "Description"
    for cell in hdr:
        set_cell_background(cell, "1E3A8A")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("kezzler_mock_dump.sql", "Complete single-file database dump (51 tables + 632 rows)."),
        ("run_ui.py", "Python script to launch the Web Chat Interface on http://127.0.0.1:8000."),
        ("seed_meaningful.py", "Script to re-generate realistic domain data across all tables."),
    ]

    for idx, (f_name, f_desc) in enumerate(data, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = f_name
        row_cells[1].text = f_desc
        if idx % 2 == 0:
            set_cell_background(row_cells[0], "F8FAFC")
            set_cell_background(row_cells[1], "F8FAFC")

    doc.add_paragraph() # Spacing

    # Section 2: Importing Database
    h2 = doc.add_heading("2. Step-by-Step Database Import Instructions", level=1)
    h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph().add_run("Choose ONE of the two methods below to import the database into your MySQL server:").bold = True

    # Method A
    h2_a = doc.add_heading("Option A: Using MySQL Workbench or DBeaver (GUI)", level=2)
    h2_a.runs[0].font.color.rgb = RGBColor(51, 65, 85)

    p_a = doc.add_paragraph()
    p_a.add_run("1. Open ").font.name = "Arial"
    p_a.add_run("MySQL Workbench").bold = True
    p_a.add_run(" or ").font.name = "Arial"
    p_a.add_run("DBeaver").bold = True
    p_a.add_run(".\n2. Connect to your local MySQL server.\n3. Click ")
    p_a.add_run("File -> Open SQL Script").bold = True
    p_a.add_run(" and select ")
    p_a.add_run("kezzler_mock_dump.sql").bold = True
    p_a.add_run(".\n4. Click the ")
    p_a.add_run("Execute (⚡)").bold = True
    p_a.add_run(" button to import all tables and data.")

    # Method B
    h2_b = doc.add_heading("Option B: Using Command Prompt / Terminal", level=2)
    h2_b.runs[0].font.color.rgb = RGBColor(51, 65, 85)

    p_b1 = doc.add_paragraph()
    p_b1.add_run("Open Command Prompt or Terminal and run the following two commands:")

    # Code Box
    table_cmd = doc.add_table(rows=1, cols=1)
    table_cmd.style = 'Table Grid'
    cell_cmd = table_cmd.rows[0].cells[0]
    set_cell_background(cell_cmd, "0F172A")
    p_cmd = cell_cmd.paragraphs[0]
    run_cmd = p_cmd.add_run(
        '# Step 1: Create empty database\n'
        'mysql -u root -p -e "CREATE DATABASE IF NOT EXISTS kezzler;"\n\n'
        '# Step 2: Import full database dump\n'
        'mysql -u root -p kezzler < kezzler_mock_dump.sql'
    )
    run_cmd.font.name = "Consolas"
    run_cmd.font.size = Pt(9.5)
    run_cmd.font.color.rgb = RGBColor(226, 232, 240)

    doc.add_paragraph() # Spacing

    # Section 3: Launching UI
    h3 = doc.add_heading("3. Running the Web Chat Interface", level=1)
    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    p_ui = doc.add_paragraph()
    p_ui.add_run("Once the database is imported, start the interactive Text-to-SQL Chat Interface:")

    table_ui = doc.add_table(rows=1, cols=1)
    table_ui.style = 'Table Grid'
    cell_ui = table_ui.rows[0].cells[0]
    set_cell_background(cell_ui, "0F172A")
    p_ui_cmd = cell_ui.paragraphs[0]
    run_ui_cmd = p_ui_cmd.add_run(
        '# 1. Start the web application\n'
        'python run_ui.py\n\n'
        '# 2. Open your web browser at:\n'
        'http://127.0.0.1:8000'
    )
    run_ui_cmd.font.name = "Consolas"
    run_ui_cmd.font.size = Pt(9.5)
    run_ui_cmd.font.color.rgb = RGBColor(226, 232, 240)

    doc.add_paragraph() # Spacing

    # Section 4: Sample Questions
    h4 = doc.add_heading("4. Sample Queries to Test", level=1)
    h4.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    p_q = doc.add_paragraph()
    p_q.add_run("Try asking these natural language questions in the Chat UI:")

    questions = [
        ("Standard Query:", "What are the active distributors?"),
        ("Revenue Query:", "Show top 5 distributors by total sales revenue excluding returned sales."),
        ("6-Table Complex Query:", "Show total revenue by zone, region, territory, distributor and product."),
        ("Ambiguous Query:", "Who is our best distributor? (Agent clarifies assumption)"),
        ("Out of Scope Query:", "What is each distributor's lifetime value (LTV)? (Agent declines)"),
        ("Security Test:", "Delete returned sales data (Blocked by safety guardrails)"),
    ]

    for category, q_text in questions:
        p_item = doc.add_paragraph(style='List Bullet')
        r_cat = p_item.add_run(f"{category} ")
        r_cat.bold = True
        r_cat.font.color.rgb = RGBColor(30, 58, 138)
        p_item.add_run(f'"{q_text}"')

    doc.add_paragraph() # Spacing

    # Footer Note
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("Kezzler AI Agent — Safe Test Environment")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(OUTPUT_DOCX)
    print(f"Generated Word document: {OUTPUT_DOCX}")

if __name__ == "__main__":
    create_guide()
